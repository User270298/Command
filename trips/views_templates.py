from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponseForbidden, HttpResponse
from django.contrib.auth import get_user_model
from django.db.models import Q
from django.utils import timezone
from django.urls import reverse
import logging
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.contrib.auth.decorators import user_passes_test
from equipment.models import EquipmentRecord
from django.db.models import Count
from datetime import datetime, timedelta
from .forms import TechnicalStaffRecordForm
from .models import TechnicalStaffRecord

from .models import BusinessTrip, Organization
from equipment.models import EquipmentRecord

import openpyxl
from openpyxl.styles import Alignment, Font, Border, Side
from openpyxl.utils import get_column_letter
from io import BytesIO

User = get_user_model()

@login_required
def trip_list_view(request):
    """
    Отображает список командировок пользователя
    """
    trips = BusinessTrip.objects.filter(members=request.user).order_by('-start_date')
    return render(request, 'trips/trip_list.html', {'trips': trips})

@login_required
def trip_detail_view(request, trip_id):
    """
    Отображает детальную информацию о командировке
    """
    trip = get_object_or_404(BusinessTrip, id=trip_id, members=request.user)
    # Получаем все организации командировки (включая закрытые)
    organizations = trip.organizations.all().order_by('-created_date')
    return render(request, 'trips/trip_detail.html', {
        'trip': trip,
        'organizations': organizations
    })

@login_required
def create_trip_view(request):
    """
    Отображает форму для создания новой командировки
    """
    # Получаем всех доступных пользователей
    available_users = BusinessTrip.get_available_users()
    
    if not available_users.exists():
        messages.warning(request, 'В данный момент нет доступных пользователей для командировки. Все пользователи уже участвуют в активных командировках.')
        return redirect('trip_list')
    
    context = {
        'available_users': available_users
    }
    
    return render(request, 'trips/create_trip.html', context)

@login_required
def trip_create_view(request):
    user = request.user
    logger = logging.getLogger(__name__)
    
    # Проверяем, не является ли пользователь уже участником или старшим в активной командировке
    active_trips = BusinessTrip.objects.filter(
        Q(senior=user) | Q(members=user),
        end_date__isnull=True
    )
    
    if active_trips.exists():
        messages.error(request, 'Вы уже участвуете в активной командировке')
        return redirect('trip_list')
    
    if request.method == 'POST':
        name = request.POST.get('name')
        city = request.POST.get('city')
        member_ids = request.POST.getlist('members')
        org_names = request.POST.getlist('org_names[]')
        
        if not name or not city:
            messages.error(request, 'Пожалуйста, заполните все обязательные поля')
            return redirect('trip_create')
        
        # Проверяем, что количество участников не превышает 3 (не считая старшего)
        if len(member_ids) > 3:
            messages.error(request, 'Максимальное количество участников в командировке - 4 человека (включая старшего)')
            return redirect('trip_create')
        
        # Проверяем, что есть хотя бы одна организация
        if not org_names or not any(org_name.strip() for org_name in org_names):
            messages.error(request, 'Пожалуйста, укажите хотя бы одну организацию')
            return redirect('trip_create')
        
        try:
            # Создаем командировку
            trip = BusinessTrip.objects.create(
                name=name,
                city=city,
                senior=user
            )
            
            # Добавляем организации
            for org_name in org_names:
                if org_name.strip():  # Проверяем, что название не пустое
                    Organization.objects.create(
                        name=org_name.strip(),
                        trip=trip
                    )
            
            # Добавляем участников
            for member_id in member_ids:
                try:
                    member = User.objects.get(id=member_id)
                    # Проверяем, что участник не в активной командировке
                    if not BusinessTrip.objects.filter(
                        Q(senior=member) | Q(members=member),
                        end_date__isnull=True
                    ).exists():
                        trip.members.add(member)
                    else:
                        messages.warning(
                            request,
                            f'Пользователь {member.get_full_name() or member.username} уже участвует в другой командировке'
                        )
                except User.DoesNotExist:
                    messages.warning(request, f'Пользователь с ID {member_id} не найден')
            
            # Добавляем старшего как участника
            trip.members.add(user)
            
            messages.success(request, 'Командировка успешно создана')
            return redirect('trip_detail', trip_id=trip.id)
            
        except Exception as e:
            messages.error(request, f'Ошибка при создании командировки: {str(e)}')
            return redirect('trip_create')
    
    # Получаем доступных пользователей для выбора
    available_users = BusinessTrip.get_available_users(exclude_user=user)
    logger.debug(f"Доступных пользователей в представлении: {available_users.count()}")
    
    # Проверяем, есть ли вообще пользователи в системе
    all_users = User.objects.all()
    logger.debug(f"Всего пользователей в системе: {all_users.count()}")
    
    # Если нет доступных пользователей, но есть пользователи в системе,
    # показываем форму без выбора участников
    if not available_users.exists() and all_users.exists():
        logger.debug("Нет доступных пользователей, но есть пользователи в системе")
        context = {
            'available_users': User.objects.exclude(id=user.id),
            'user': user,
            'no_available_users': True
        }
    else:
        context = {
            'available_users': available_users,
            'user': user,
            'no_available_users': False
        }
    
    return render(request, 'trips/trip_create.html', context)

@login_required
def close_trip(request, pk):
    trip = get_object_or_404(BusinessTrip, id=pk)
    
    # Проверяем, что пользователь является старшим в этой командировке
    if request.user != trip.senior:
        messages.error(request, "Только старший командировки может завершить её")
        return redirect('trip_list')
    
    # Проверяем, что командировка активна
    if trip.end_date is not None:
        messages.warning(request, "Эта командировка уже завершена")
        return redirect('trip_list')
    
    # Завершаем командировку, устанавливая дату окончания
    trip.end_date = timezone.now().date()
    trip.save()
    
    messages.success(request, f"Командировка '{trip.name}' успешно завершена")
    return redirect('trip_list')

@login_required
def organization_detail_view(request, org_id):
    user = request.user
    organization = get_object_or_404(Organization, id=org_id)
    trip = organization.trip
    if not (user == trip.senior or user in trip.members.all()):
        return HttpResponseForbidden("Вы не имеете доступа к этой организации")

    # Всегда показываем все приборы этой организации
    equipment_records = organization.equipment_records.select_related('measurement_type', 'user', 'organization')

    # Статистика по типам измерений
    measurement_stats = equipment_records.values('measurement_type__name').distinct()

    is_senior = user == trip.senior
    return render(request, 'trips/organization_detail.html', {
        'organization': organization,
        'equipment_records': equipment_records,
        'measurement_stats': measurement_stats,
        'is_senior': is_senior,
    })

@login_required
def organization_create_view(request):
    user = request.user
    selected_trip_id = request.GET.get('trip')
    selected_trip = None
    
    if selected_trip_id:
        try:
            selected_trip = BusinessTrip.objects.get(id=selected_trip_id)
            # Check if user is the senior
            if user != selected_trip.senior:
                messages.error(request, 'Только старший группы может добавлять организации')
                return redirect('trip_detail', trip_id=selected_trip.id)
        except BusinessTrip.DoesNotExist:
            pass
    
    if request.method == 'POST':
        name = request.POST.get('name')
        trip_id = request.POST.get('trip')
        
        if not name or not trip_id:
            messages.error(request, 'Пожалуйста, заполните все поля')
            return redirect('organization_create')
        
        try:
            trip = BusinessTrip.objects.get(id=trip_id)
            
            # Only senior can add organizations
            if user != trip.senior:
                messages.error(request, 'Только старший группы может добавлять организации')
                return redirect('trip_detail', trip_id=trip.id)
            
            # Create organization
            organization = Organization.objects.create(
                name=name,
                trip=trip
            )
            
            messages.success(request, 'Организация успешно добавлена')
            return redirect('organization_detail', org_id=organization.id)
            
        except BusinessTrip.DoesNotExist:
            messages.error(request, 'Командировка не найдена')
            return redirect('trip_list')
    
    # Get active trips where user is senior
    trips = BusinessTrip.objects.filter(senior=user, end_date=None)
    
    context = {
        'trips': trips,
        'selected_trip': selected_trip
    }
    
    return render(request, 'trips/organization_create.html', context)

@login_required
def close_organization_view(request, org_id):
    user = request.user
    organization = get_object_or_404(Organization, id=org_id)
    
    # Only the trip's senior can close an organization
    if user != organization.trip.senior:
        return HttpResponseForbidden("Только старший группы может закрыть организацию")
    
    if request.method == 'POST':
        # This will create an Excel report and mark the organization as closed
        
        # Отметить организацию как закрытую
        organization.is_closed = True
        organization.save()
        
        messages.success(request, 'Организация успешно закрыта')
        
        # Перенаправить на страницу деталей командировки
        return redirect('trip_detail', trip_id=organization.trip.id)
    
    return redirect('organization_detail', org_id=organization.id)

@login_required
def generate_weekly_report(request, org_id):
    user = request.user
    organization = get_object_or_404(Organization, id=org_id)
    
    # Check if user is the senior of the organization's trip
    if user != organization.trip.senior:
        return HttpResponseForbidden("Только старший группы может выгружать отчеты")
    
    # Get records from the current week (Monday to Sunday)
    today = timezone.now().date()
    monday = today - timezone.timedelta(days=today.weekday())
    sunday = monday + timezone.timedelta(days=6)
    
    records = organization.equipment_records.filter(
        date_created__date__gte=monday,
        date_created__date__lte=sunday
    ).select_related('measurement_type', 'user')
    
    # Group records by user and measurement type
    report_data = {}
    for record in records:
        user_name = record.user.get_full_name() or record.user.username
        if user_name not in report_data:
            report_data[user_name] = {}
        if record.measurement_type.name not in report_data[user_name]:
            report_data[user_name][record.measurement_type.name] = 0
        report_data[user_name][record.measurement_type.name] += record.quantity
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    
    doc.add_heading(f'Отчет по организации {organization.name} за неделю', 0)
    doc.add_paragraph(f'Период: с {monday.strftime("%d.%m.%Y")} по {sunday.strftime("%d.%m.%Y")}')
    
    # Add user statistics
    for user_name, measurements in report_data.items():
        doc.add_paragraph(f'Работник (поверитель): {user_name}')
        for measurement_type, count in measurements.items():
            doc.add_paragraph(f'{count} СИ {measurement_type}', style='List Bullet')
        doc.add_paragraph()  # Add empty line between users
    
    # Save the document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=weekly_report_{organization.name}.docx'
    doc.save(response)
    return response

@login_required
def download_organization_list(request, org_id):
    user = request.user
    organization = get_object_or_404(Organization, id=org_id)
    if user != organization.trip.senior:
        return HttpResponseForbidden("Только старший группы может скачивать перечень организации")
    
    # Получить все записи
    records = organization.equipment_records.all().select_related('measurement_type', 'user')
    
    # Определяем порядок типов измерений
    measurement_order = [
        'измерения геометрических величин',
        'измерения механических величин',
        'измерения параметров потока, расхода, уровня и объема веществ',
        'измерения давления и вакуумные',
        'измерения физико-химического состава и свойств веществ',
        'теплофизические и температурные измерения',
        'измерения времени и частоты',
        'измерения электрических и магнитных величин',
        'радиотехнические и радиоэлектронные измерения',
        'измерения акустических величин',
        'оптико-физические измерения',
        'измерения характеристик ионизирующих излучений и ядерных констант',
        'метеорологические измерения',
        'специальные измерения (в соответствии со специализацией)'
    ]
    
    doc = Document()
    
    # Устанавливаем меньший шрифт для всего документа
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(5)  # Уменьшили с 8 до 7
    style.font.italic = True
    
    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('ЭТАЛОНЫ (СИ), АТТЕСТОВАННЫЕ (ПОВЕРЕННЫЕ) СОГЛАСНО ПЛАНА-ГРАФИКА\nПРЕДСТАВЛЕНИЯ ВОИНСКИМИ ЧАСТЯМИ (ПОДРАЗДЕЛЕНИЯМИ) НА АТТЕСТАЦИЮ\nЭТАЛОНОВ (ПОВЕРКУ СИ) В ПЛАНИРУЕМОМ ГОДУ')
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(5)  # Уменьшили с 8 до 7
    title_run.bold = True
    title_run.italic = True
    
    # Создаем таблицу
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    
    headers = ['№ п/п', 'Наименование', 'Тип', 'Заводской номер', 'Количество', 
              'Результат аттестации (поверки)', 'Дата выполнения аттестации (поверки)',
              'Тип ВВСТ, в состав которого входит эталон (СИ)',
              'Примечание (номер извещения о непригодности к применению (свидетельства об аттестации (о поверке))']
    
    # Настройка заголовков таблицы
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        # Стиль для заголовка
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(7)  # Уменьшили с 8 до 7
                run.bold = True
                run.italic = True
            # Все заголовки по центру
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Группируем записи по типам измерений
    measurement_groups = {}
    for record in records:
        if record.measurement_type.name not in measurement_groups:
            measurement_groups[record.measurement_type.name] = []
        measurement_groups[record.measurement_type.name].append(record)
    
    row_num = 1
    
    # Сортируем типы измерений согласно заданному порядку
    sorted_measurement_types = []
    for measurement_type in measurement_order:
        if measurement_type in measurement_groups:
            sorted_measurement_types.append(measurement_type)
    
    # Добавляем остальные типы измерений, которых нет в списке
    for measurement_type in measurement_groups.keys():
        if measurement_type not in sorted_measurement_types:
            sorted_measurement_types.append(measurement_type)
    
    # Заполняем таблицу в правильном порядке
    for measurement_type in sorted_measurement_types:
        records_list = measurement_groups[measurement_type]
        
        # Добавляем заголовок группы
        row_cells = table.add_row().cells
        row_cells[0].text = f'Средства {measurement_type}'
        row_cells[0].merge(row_cells[8])
        cell = row_cells[0]
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(7)  # Уменьшили с 8 до 7
                run.bold = True
                run.italic = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем записи для этого типа измерения
        for record in records_list:
            row_cells = table.add_row().cells
            row_cells[0].text = str(row_num)
            row_cells[1].text = record.name
            row_cells[2].text = record.device_type
            row_cells[3].text = record.serial_number or ''
            row_cells[4].text = str(record.quantity)
            row_cells[5].text = record.status
            row_cells[6].text = f'С {timezone.now().strftime("%d.%m.%Y")} по {(timezone.now() + timezone.timedelta(days=365)).strftime("%d.%m.%Y")}'
            row_cells[7].text = record.tech_name
            row_cells[8].text = record.notes or ''
            
            # Настройка стиля для всех ячеек
            for i, cell in enumerate(row_cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(7)  # Уменьшили с 8 до 7
                        run.italic = True
                    # Все ячейки по центру
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row_num += 1

    # Подсчет количества приборов
    all_records = organization.equipment_records.all().select_related('measurement_type', 'user')
    total_devices = sum(r.quantity for r in all_records)
    rejected_devices = sum(r.quantity for r in all_records if str(r.status).strip().lower() in ['брак', 'негоден', 'не годен', 'забраковано'])
    
    # Добавляем итоговую информацию
    summary_paragraph = doc.add_paragraph()
    summary_text = f"Всего приборов - {total_devices}, из них забраковано - {rejected_devices}"
    run = summary_paragraph.add_run(summary_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(7)  # Уменьшили с 8 до 7
    run.italic = True
    summary_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=perechen_{organization.name}.docx'
    doc.save(response)
    return response

@login_required
def complete_organization(request, org_id):
    user = request.user
    organization = get_object_or_404(Organization, id=org_id)
    if user != organization.trip.senior:
        return HttpResponseForbidden("Только старший группы может завершать организации")
    if request.method == 'POST':
        organization.is_closed = True
        organization.save()
        messages.success(request, 'Организация успешно закрыта')
        return redirect('organization_detail', org_id=organization.id)
    return HttpResponseForbidden("Invalid request method")

@login_required
def download_plan_assignment_excel(request, org_id):
    user = request.user
    organization = get_object_or_404(Organization, id=org_id)
    if user != organization.trip.senior:
        return HttpResponseForbidden("Только старший группы может скачивать План-Задание")
    records = organization.equipment_records.all().select_related('measurement_type', 'user')

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "План-Задание"

    # Шапка
    ws.merge_cells('A1:K1')
    ws['A1'] = 'I. Аттестация эталонов и поверка средств измерений подразделений воинских частей (организаций)'
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws['A1'].font = Font(bold=True)

    ws.merge_cells('A2:E2')
    ws['A2'] = 'Задание (аттестовать, поверить, отремонтировать)'
    ws['A2'].alignment = Alignment(horizontal='center', vertical='center')
    ws.merge_cells('F2:K2')
    ws['F2'] = 'Отчет'
    ws['F2'].alignment = Alignment(horizontal='center', vertical='center')

    # Заголовки
    headers = [
        '№ п/п',
        'Наименование эталонов (СИ)',
        'Тип эталона (СИ)',
        'Заводской номер',
        'Количество',
        'Планируемое время аттестации (поверки), ч',
        'Результат аттестации (поверки) (годен, брак)',
        'Затрачено времени на аттестацию (поверку), ч',
        'Израсходованные ЗИП, материалы, ГСМ при аттестации (поверке), ремонте',
        'Фамилия и инициалы лица, проводившего аттестацию (поверку), ремонт',
        'Примечание',
    ]
    ws.append(headers)

    # Стили для заголовков
    for col in range(1, len(headers) + 1):
        cell = ws.cell(row=3, column=col)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        ws.column_dimensions[get_column_letter(col)].width = 18

    # Данные
    for idx, record in enumerate(records, 1):
        ws.append([
            idx,
            record.name,
            record.device_type,
            record.serial_number or '',
            record.quantity,
            '',  # Планируемое время аттестации (поверки), ч
            record.status,
            '',  # Затрачено времени на аттестацию (поверку), ч
            '',  # Израсходованные ЗИП, материалы, ГСМ
            record.user.get_full_name() or record.user.username,
            record.notes or '',
        ])
        for col in range(1, len(headers) + 1):
            ws.cell(row=idx+3, column=col).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    # Границы
    thin = Side(border_style="thin", color="000000")
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=len(headers)):
        for cell in row:
            cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)

    # Ответ
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    response = HttpResponse(output.read(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename=plan_assignment_{organization.name}.xlsx'
    return response

def get_current_week_start(now):
    # Всегда возвращать ближайшую прошедшую субботу 16:00
    weekday = now.weekday()  # 0=Monday, 5=Saturday
    days_since_saturday = (weekday - 5) % 7
    last_saturday = now - timedelta(days=days_since_saturday)
    last_saturday_16 = last_saturday.replace(hour=16, minute=0, second=0, microsecond=0)
    if now < last_saturday_16:
        last_saturday_16 -= timedelta(days=7)
    return last_saturday_16

# В местах, где вычисляется week_end:
# week_end = week_start + timedelta(days=7)

@user_passes_test(lambda u: u.is_authenticated and getattr(u, 'isAdmin', False))
def admin_stats_view(request):
    from collections import defaultdict
    from equipment.models import EquipmentRecord, MeasurementType
    from trips.models import BusinessTrip, Organization
    from django.contrib.auth import get_user_model
    from .models import TechnicalStaffRecord

    # Получаем все командировки
    trips = BusinessTrip.objects.prefetch_related('organizations', 'organizations__equipment_records', 'members', 'senior')

    # Общий сводный отчет по всем командировкам
    all_stats = defaultdict(lambda: defaultdict(int))  # {user: {measurement_type: count}}
    for record in EquipmentRecord.objects.all():
        all_stats[record.user.username][record.measurement_type.name] += record.quantity

    # Считаем общее количество поверенных СИ для каждого пользователя
    all_stats_totals = {user: sum(mtypes.values()) for user, mtypes in all_stats.items()}

    # Получаем текущую неделю (понедельник 00:00 - воскресенье 23:59)
    now = timezone.now()
    week_start = get_current_week_start(now)
    week_end = week_start + timedelta(days=7)

    # --- ДОБАВЛЯЕМ: выборка числа тех. состава за неделю по организациям ---
    tech_staff_records = TechnicalStaffRecord.objects.filter(created_at__gte=week_start, created_at__lt=week_end)
    tech_staff_by_org = {rec.organization_id: rec for rec in tech_staff_records}

    # Общий сводный отчет по всем командировкам ЗА ТЕКУЩУЮ НЕДЕЛЮ
    all_stats_week = defaultdict(lambda: defaultdict(int))  # {user: {measurement_type: count}}
    for record in EquipmentRecord.objects.filter(date_created__gte=week_start, date_created__lt=week_end):
        all_stats_week[record.user.username][record.measurement_type.name] += record.quantity

    all_stats_week_totals = {user: sum(mtypes.values()) for user, mtypes in all_stats_week.items()}

    # Структура для подробного отчета по каждой командировке
    trip_stats = []
    for trip in trips:
        trip_info = {
            'name': trip.name,
            'city': trip.city,
            'senior': trip.senior,
            'organizations': [],
        }
        for org in trip.organizations.all():
            org_info = {
                'id': org.id,
                'name': org.name,
                'is_closed': org.is_closed,
                'workers': defaultdict(lambda: defaultdict(int)),
                'records': [],
                'tech_staff': tech_staff_by_org.get(org.id),
                'usernames_to_names': {},
            }
            records = org.equipment_records.filter(
                date_created__gte=week_start,
                date_created__lt=week_end
            ).select_related('user', 'measurement_type')
            for rec in records:
                org_info['workers'][rec.user.username][rec.measurement_type.name] += rec.quantity
                org_info['records'].append(rec)
                org_info['usernames_to_names'][rec.user.username] = rec.user.get_full_name() or rec.user.username
            week_stats_exist = any(sum(mtypes.values()) > 0 for mtypes in org_info['workers'].values())
            org_info['workers'] = {k: dict(v) for k, v in org_info['workers'].items()}
            org_info['has_week_stats'] = week_stats_exist
            # Краткая статистика по работникам и типам измерений
            short_stats = defaultdict(lambda: defaultdict(lambda: {'total': 0, 'rejected': 0}))
            for rec in org_info['records']:
                worker = rec.user.get_full_name() or rec.user.username if rec.user else "—"
                mtype_name = rec.measurement_type.name if rec.measurement_type else "Неизвестно"
                short_stats[worker][mtype_name]['total'] += rec.quantity
                if str(rec.status).strip().lower() in ['брак', 'негоден', 'не годен', 'забраковано']:
                    short_stats[worker][mtype_name]['rejected'] += rec.quantity
            # Исправление: удалить все ключи, значения которых не словарь
            bad_keys = [k for k, v in short_stats.items() if not isinstance(v, dict)]
            for k in bad_keys:
                del short_stats[k]
            # Преобразуем к обычному dict для шаблона
            org_info['short_stats'] = {k: {mt: dict(vv) for mt, vv in v.items()} for k, v in short_stats.items() if isinstance(v, dict)}
            trip_info['organizations'].append(org_info)
        trip_stats.append(trip_info)
    return render(request, 'admin_stats.html', {
        'all_stats': all_stats,
        'all_stats_totals': all_stats_totals,
        'all_stats_week': all_stats_week,
        'all_stats_week_totals': all_stats_week_totals,
        'trip_stats': trip_stats,
        'week_start': week_start,
        'week_end': week_end,
    }) 

@login_required
def tech_staff_entry_view(request):
    user = request.user
    # Проверяем, что пользователь — старший группы
    if not getattr(user, 'is_senior', False):
        return HttpResponseForbidden("Только старший группы может вносить данные тех. состава.")

    now = timezone.now()
    # Проверка: только до 12:00 пятницы
    if not (now.weekday() == 4 and now.hour < 12):
        messages.error(request, "Вносить данные можно только до 12:00 пятницы.")
        return redirect('trip_list')

    # Получаем все незакрытые организации, по которым нет записи за эту неделю
    monday = now.date() - timedelta(days=now.weekday())
    organizations = Organization.objects.filter(is_closed=False, trip__senior=user)
    orgs_to_fill = []
    for org in organizations:
        if not TechnicalStaffRecord.objects.filter(organization=org, week=monday).exists():
            orgs_to_fill.append(org)

    forms_list = []
    if request.method == 'POST':
        success = True
        for org in orgs_to_fill:
            value = request.POST.get(f'value_{org.id}')
            if value:
                form = TechnicalStaffRecordForm({'organization': org.id, 'value': value})
                if form.is_valid():
                    rec = form.save(commit=False)
                    rec.user = user
                    rec.week = monday
                    rec.save()
                else:
                    success = False
                    forms_list.append((org, form))
            else:
                forms_list.append((org, TechnicalStaffRecordForm(organization=org)))
        if success:
            messages.success(request, "Данные успешно внесены!")
            return redirect('trip_list')
    else:
        for org in orgs_to_fill:
            forms_list.append((org, TechnicalStaffRecordForm(organization=org)))

    return render(request, 'trips/tech_staff_entry.html', {'forms_list': forms_list}) 